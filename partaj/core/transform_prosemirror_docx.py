"""
Class to transform rich text view front component to docx format text
"""

import json

from django.utils import dateformat

import docx


class TransformProsemirrorDocx:
    """
    Transfom objects with properties that contain ProseMirror-formatted text to a docx document.
    """

    list_type = {"bullet_list": "List Bullet", "ordered_list": "List Number"}

    def __init__(self):
        self.doc_word = docx.Document()

    def referral_to_docx(self, referral):
        """
        Transform Referral into a docx document
        """
        self.doc_word = docx.Document()

        self.doc_word.add_heading("Saisine N°" + str(referral.id), 0)
        self.doc_word.add_paragraph(
            "créée le " + dateformat.format(referral.created_at, "j F Y"),
        )

        self.doc_word.add_heading("Demandeurs", level=1)
        paragraph_users = self.doc_word.add_paragraph()
        paragraph_users.paragraph_format.line_spacing = 1
        for user in referral.users.all():
            paragraph_users.add_run(
                user.get_full_name() + " - " + user.unit_name + str("\n")
            )

        self.doc_word.add_heading("Thème de la saisine:", level=1)
        self.doc_word.add_paragraph(referral.topic.name)

        self.doc_word.add_heading("Objet:", level=1)
        self.doc_word.add_paragraph(referral.object)

        self.doc_word.add_heading("Question:", level=1)
        self.transform_richtext(referral.question)

        self.doc_word.add_heading("Contexte:", level=1)
        self.transform_richtext(referral.context)

        self.doc_word.add_heading("Travaux préalables:", level=1)
        self.transform_richtext(referral.prior_work)

        self.doc_word.add_heading("Délai de réponse attendu:", level=1)
        self.doc_word.add_paragraph(referral.urgency_level.name)

        self.doc_word.add_heading("Justification de l'urgence:", level=1)
        self.doc_word.add_paragraph(referral.urgency_explanation)

        return self.doc_word

    def transform_richtext(self, text):
        """
        Transform text from prosemirror format to docx format.
        """
        data = json.loads(text)
        for paragraph in data["doc"]["content"]:
            if paragraph["type"] == "paragraph":
                new_paragraph = self.doc_word.add_paragraph()  # add new paragraph
                new_paragraph.paragraph_format.line_spacing = 1
                if "content" in paragraph:
                    for content in paragraph["content"]:
                        if content["type"] == "text":
                            self.transform_text(content, new_paragraph)

            if paragraph["type"] in self.list_type:
                self.transform_list(paragraph["content"], paragraph["type"], 1)

            if paragraph["type"] == "blockquote":
                self.transform_blockquote(paragraph["content"])

    def transform_blockquote(self, blockquote_list):
        """
        Transform recursively blockquote to docx blockquote format.
        """
        for blockquote_item in blockquote_list:
            if blockquote_item["type"] == "paragraph":
                blockquote_paragraph = self.doc_word.add_paragraph(
                    style="Intense Quote"
                )
                blockquote_paragraph.paragraph_format.line_spacing = 1
                blockquote_paragraph.paragraph_format.space_after = 1
                if "content" in blockquote_item:
                    for content in blockquote_item["content"]:
                        if content["type"] == "text":
                            self.transform_text(content, blockquote_paragraph)

            if blockquote_item["type"] == "blockquote":  # if text  => transform text
                self.transform_blockquote(blockquote_item["content"])

            if blockquote_item["type"] in self.list_type:  # if list   => transform list
                # add new paragraph
                self.transform_list(
                    blockquote_item["content"],
                    blockquote_item["type"],
                    1,
                )

    def transform_text(self, text, new_paragraph):
        """
        Transform text to docx text format.
        """
        runner = new_paragraph.add_run(text["text"])

        if "marks" in text:
            for mark in text["marks"]:
                if mark["type"] == "strong":
                    runner.font.bold = True
                if mark["type"] == "em":
                    runner.font.italic = True
                if mark["type"] == "underline":
                    runner.font.underline = True

    def transform_list(self, list_to_transform, type_list, level):
        """
        Transform recursively list to docx list format.
        """
        for list_item in list_to_transform:
            for list_item_content in list_item["content"]:
                if (
                    list_item_content["type"] in self.list_type
                ):  # if list  => transform list
                    self.transform_list(
                        list_item_content["content"],
                        list_item_content["type"],
                        level + 1,
                    )
                if (
                    list_item_content["type"] == "paragraph"
                ):  # if text  => transform text
                    # first list item doesn't have number, others yes
                    style = (
                        lambda: self.list_type[type_list],
                        lambda: self.list_type[type_list] + " " + str(level),
                    )[level > 1]()
                    list_paragraph = self.doc_word.add_paragraph(style=style)
                    list_paragraph.paragraph_format.line_spacing = 1
                    list_paragraph.paragraph_format.space_after = 1
                    if "content" in list_item_content:
                        for item_text in list_item_content["content"]:
                            self.transform_text(item_text, list_paragraph)
