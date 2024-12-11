"""
Class to transform rich text view front component to docx format text
"""

from django.utils import dateformat

import docx

from partaj.core.models import RequesterUnitType


class ReferralDocx:
    """
    Transfom referral to docx.
    """

    list_type = {"bullet_list": "List Bullet", "ordered_list": "List Number"}

    def __init__(self):
        self.doc_word = docx.Document()

    def referral_to_docx(self, referral):
        """
        Transform Referral into a docx document
        """
        self.doc_word = docx.Document()

        self.doc_word.add_heading("Saisine N°" + str(referral.id), 0)
        self.doc_word.add_paragraph(
            "transmise le " + dateformat.format(referral.sent_at, "j F Y"),
        )

        self.doc_word.add_heading("Demandeurs", level=1)
        paragraph_users = self.doc_word.add_paragraph()
        paragraph_users.paragraph_format.line_spacing = 1
        for user in referral.users.all():
            paragraph_users.add_run(
                user.get_full_name() + " - " + user.unit_name + str("\n")
            )

        self.doc_word.add_heading("Thème de la saisine:", level=1)
        self.doc_word.add_paragraph(referral.topic.name)

        self.doc_word.add_heading("Objet:", level=1)
        self.doc_word.add_paragraph(referral.object)

        self.doc_word.add_heading("Question:", level=1)
        self.doc_word.add_paragraph(referral.question)

        self.doc_word.add_heading("Contexte:", level=1)
        self.doc_word.add_paragraph(referral.context)

        self.doc_word.add_heading("Travaux préalables:", level=1)

        if referral.has_prior_work:
            self.doc_word.add_paragraph("Oui")
            if referral.requester_unit_type == RequesterUnitType.DECENTRALISED_UNIT:
                self.doc_word.add_heading("Contact :", level=2)
                self.doc_word.add_paragraph(referral.requester_unit_contact)
            self.doc_word.add_paragraph(referral.prior_work)
        else:
            self.doc_word.add_paragraph("Non")
            if referral.requester_unit_type == RequesterUnitType.DECENTRALISED_UNIT:
                self.doc_word.add_heading("Justification:", level=2)
                self.doc_word.add_paragraph(referral.no_prior_work_justification)

        self.doc_word.add_heading("Délai de réponse attendu:", level=1)
        self.doc_word.add_paragraph(referral.urgency_level.name)

        self.doc_word.add_heading("Justification de l'urgence:", level=1)
        self.doc_word.add_paragraph(referral.urgency_explanation)

        return self.doc_word
